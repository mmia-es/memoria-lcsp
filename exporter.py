from docx import Document

def export_basic(doc_data: dict, out_path: str):
    doc = Document()
    doc.add_heading('Memoria justificativa (borrador)', level=1)
    for key, text in doc_data.items():
        doc.add_heading(key, level=2)
        for para in text.split('\n'):
            doc.add_paragraph(para)
    doc.save(out_path)

def export_with_template(template_path: str, doc_data: dict, out_path: str):
    doc = Document(template_path)

    def replace_in_paragraph(paragraph, mapping):
        for k, v in mapping.items():
            target = '{{' + k + '}}'
            if target in paragraph.text:
                inline = paragraph.runs
                text = paragraph.text.replace(target, v)
                for i in range(len(inline)-1, -1, -1):
                    paragraph.runs[i]._r.getparent().remove(paragraph.runs[i]._r)
                paragraph.add_run(text)

    def replace_in_table(table, mapping):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, mapping)

    mapping = {k: (v or '') for k, v in doc_data.items()}
    for p in doc.paragraphs:
        replace_in_paragraph(p, mapping)
    for tbl in doc.tables:
        replace_in_table(tbl, mapping)

    doc.save(out_path)
